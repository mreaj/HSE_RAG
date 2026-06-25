"""
Turn raw file bytes / HTML into clean text, then into overlapping chunks.

Extraction:
  pdf  -> PyMuPDF (fitz), page by page
  docx -> python-docx, paragraph text
  html -> trafilatura (falls back to a crude tag strip)
  txt/md -> utf-8 decode

Chunking is word-based with overlap (config.CHUNK_SIZE / CHUNK_OVERLAP).
Each chunk keeps an optional page number for citation.
"""
from __future__ import annotations

import re
from typing import Optional

from core.config import get_settings

_settings = get_settings()


# ── extraction ───────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, ext: str) -> list[tuple[str, Optional[int]]]:
    """Return list of (text, page_number). page_number is None for non-paged formats."""
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return [(_extract_docx(file_bytes), None)]
    if ext in ("html", "htm"):
        return [(extract_html(file_bytes.decode("utf-8", errors="replace")), None)]
    # txt / md / anything text-like
    return [(file_bytes.decode("utf-8", errors="replace"), None)]


def _extract_pdf(file_bytes: bytes) -> list[tuple[str, Optional[int]]]:
    import fitz  # PyMuPDF
    out: list[tuple[str, Optional[int]]] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for i, page in enumerate(doc, start=1):
            txt = page.get_text("text") or ""
            if txt.strip():
                out.append((txt, i))
    return out


def _extract_docx(file_bytes: bytes) -> str:
    import io
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # include table cell text too
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_html(html: str, url: Optional[str] = None) -> str:
    try:
        import trafilatura
        text = trafilatura.extract(
            html, include_tables=True, favor_precision=True, url=url
        )
        if text:
            return text
    except Exception:
        pass
    clean = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html)
    clean = re.sub(r"<[^>]+>", " ", clean)
    return re.sub(r"\s+", " ", clean).strip()


# ── chunking ─────────────────────────────────────────────────────────────────

def chunk_text(text: str, page: Optional[int] = None) -> list[dict]:
    words = text.split()
    if not words:
        return []
    size = max(50, _settings.chunk_size)
    overlap = min(_settings.chunk_overlap, size - 1)
    step = size - overlap
    chunks: list[dict] = []
    for start in range(0, len(words), step):
        piece = " ".join(words[start:start + size]).strip()
        if piece:
            chunks.append({"text": piece, "page_number": page})
        if start + size >= len(words):
            break
    return chunks


def pages_to_chunks(pages: list[tuple[str, Optional[int]]]) -> list[dict]:
    out: list[dict] = []
    for text, page in pages:
        out.extend(chunk_text(text, page))
    return out
